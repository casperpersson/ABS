import docx

d = docx.Document()
d.add_paragraph('Hello this is a paragraph ')
d.add_paragraph('this is paragraph 2')
p = d.paragraphs[0]
p.add_run('this is a new run')
p.runs[1].bold = True

d.save('demo2.docx')
